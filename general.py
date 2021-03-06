# list all keywords
def list_keywords():
    import keyword
    return keyword.kwlist


# show variable's unique identity
def show_var_id(var):
    id(var)


# show object's type
def show_obj_type(var):
    type(var)


# list all standard types
def list_standard_types():
    import types
    dir(types)


# calculate factorial of N
def factorial(n):
    assert n > 0
    if n != 0:
        return n * factorial(n-1)
    else:
        return 1


# list all attributes in a module
def list_all_module():
    import copy
    return [name for name in dir(copy) if name[0] != '_']


# use iterator constructor to iterate over a list
def iter_constructor():
    iter_list = [1, 2, 3]
    iterator = iter(iter_list)
    print(list(iterator))


# example of a generator
def generator_ex(low, high):
    for a in range(low, high):
        yield a
    for b in range(high-2, low-1, -1):
        yield b


# example of a class
class Range(object):
    def __init__(self, start=0, end=10):
        self.counter = start
        self.end = end

    def next(self):
        if self.counter < self.end:
            res = self.counter
            self.counter += 1
            return res
        else:
            return None


# display Zen of Python
def zen():
    import this


# calculate square root using Newton's method
def square_root(a, estimated):
    epsilon = 0.0000001
    while True:
        print(estimated)
        result = (estimated + a / estimated) / 2
        if result == estimated:
            break
        elif abs(result - estimated) < epsilon:
            break
        estimated = result


# find all occurrences of emails in a text block from clipboard
def find_all_emails(regex):
    import pyperclip

    text = str(pyperclip.paste())
    matches = []
    for groups in regex.findall(text):
        matches.append(groups[0])

    if len(matches) > 0:
        pyperclip.copy('\n'.join(matches))
        print('Copied:')
        print('\n'.join(matches))
    else:
        print('Nothing found.')


# copy files, returns a path to the newly copied file
def copy_file(source, destination):
    import shutil, os
    os.chdir('C:\\')
    return shutil.copy(source, destination)


# list files and write them to a text file
def list_files(path, extension):
    import zipfile, os
    for folder, subfolders, filenames in os.walk(path):
        for filename in filenames:
            if filename.endswith('.{}'.format(extension)):
                print("FOLDER: {}".format(folder))
                print(filename)
                with open("test.txt", 'a', encoding='utf-8') as log:
                    log.write(folder + '\n')
                    log.write(u"{}\n".format(filename))
            # elif filename.endswith('.zip'):
                # with zipfile.ZipFile(filename) as archive:
                    # print(archive.namelist())


# open a web-browser
def web_browser(url):
    import webbrowser
    webbrowser.open(url)


# download a web page
def dl_web_page(url):
    import requests
    res = requests.get(url)
    try:
        res.raise_for_status()
    except Exception as exc:
        print('There was a problem handling your request: {}'.format(exc))
    print('Status code: {}'.format(res.status_code))
    print(res.text[:250])
    return res


# write the web page to a text file
def write_web_to_txt(res):
    dl_file = open('web.txt', 'wb')
    for chunk in res.iter_content(100000):
        dl_file.write(chunk)
    dl_file.close()


# parse html
def parse_html(res=None, filename=''):
    import bs4
    if res is None and filename == '':
        return
    elif res is None and filename != '':
        file = open(filename)
        soup = bs4.BeautifulSoup(file, "lxml")
        print(soup.text[:250])
    else:
        try:
            res.raise_for_status()
        except Exception as exc:
            print('There was a problem handling your request: {}'.format(exc))
        print('Status code: {}'.format(res.status_code))
        soup = bs4.BeautifulSoup(res.text, "lxml")
        print(soup.text[:250])


# Selenium test
def selenium_ex(url, classname, text, email, passw):
    from selenium import webdriver
    browser = webdriver.Firefox()
    browser.get(url)

    try:
        elem = browser.find_element_by_class_name(classname)
        print('found {}'.format(elem.tag_name))
    except:
        print('did not find element specified')
    linkelem = browser.find_element_by_link_text(text)
    linkelem.click()

    emailelem = browser.find_element_by_id('login-username')
    emailelem.send_keys(email)
    passwordelem = browser.find_element_by_id('login-passwd')
    passwordelem.send_keys(passw)
    passwordelem.submit()


# parse xlsx files
def parse_excel(file):
    import openpyxl
    wb = openpyxl.load_workbook(file)
    sheet = wb.get_sheet_by_name('Sheet1')
    for i in range(1, sheet.max_row + 1):
        print('{} - {}'.format(i, sheet.cell(row=i, column=2).value))


def list_excel_rows(file):
    import openpyxl
    wb = openpyxl.load_workbook(file)
    sheet = wb.get_sheet_by_name('Sheet1')
    for rowObj in sheet['A1':'C3']:
        for cellObj in rowObj:
            print('{} - {}'.format(cellObj.coordinate, cellObj.value))
        print('--- end of row ---')


def list_excel_columns(file):
    import openpyxl
    wb = openpyxl.load_workbook(file)
    sheet = wb.get_sheet_by_name('Sheet1')
    for cellObj in sheet.columns[1]:
        print(cellObj.value)


def print_out_excel_data(file):
    import openpyxl, pprint
    print('Opening workbook...')
    wb = openpyxl.load_workbook(file)
    sheet = wb.get_sheet_by_name('Population by Census Tract')
    countyData = {}

    for row in range(2, sheet.max_row + 1):
        state = sheet['B' + str(row)].value
        county = sheet['C' + str(row)].value
        pop = sheet['D' + str(row)].value

        countyData.setdefault(state, {})
        countyData[state].setdefault(county, {'tracts': 0, 'pop': 0})
        countyData[state][county]['tracts'] += 1
        countyData[state][county]['pop'] += int(pop)

    print('Writing results...')
    resultFile = open('census.py', 'w')
    resultFile.write('allData = ' + pprint.pformat(countyData))
    resultFile.close()
    print('Done.')


# create xlsx files
def update_excel_data(file):
    import openpyxl
    wb = openpyxl.load_workbook(file)
    sheet = wb.get_sheet_by_name('Sheet')

    PRICE_UPDATES = {'Garlic': 3.07,
                     'Celery': 1.19,
                     'Lemon': 1.27}

    for rowNum in range(2, sheet.max_row):
        produceName = sheet.cell(row=rowNum, column=1).value
        if produceName in PRICE_UPDATES:
            sheet.cell(row=rowNum, column=2).value = PRICE_UPDATES[produceName]

    wb.save('prod.xlsx')


# parse pdf files
def parse_pdf(file):
    import PyPDF2
    pdffo = open(file, 'rb')
    pdfreader = PyPDF2.PdfFileReader(pdffo)
    print(pdfreader.isEncrypted)
    pageobj = pdfreader.getPage(0)
    print(pageobj.extractText())


# parse docx files
def parse_word(file):
    import docx
    doc = docx.Document(file)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


# create docx file
def create_word(file):
    import docx
    doc = docx.Document()
    doc.add_heading('Slim Shady', 0)
    doc.add_paragraph('Hi, ')
    paraobj = doc.add_paragraph('My name is ')
    paraobj.add_run('Slim Shady.')
    doc.save(file)


# parsing csv files
def parse_csv(file):
    import csv
    fileobj = open(file)
    reader = csv.reader(fileobj)
    for row in reader:
        print('Row # {}: {}'.format(str(reader.line_num), str(row)))


# create csv file
def create_csv(file):
    import csv
    output = open(file, 'w', newline='')  # newline is needed for Windows only (removes the extra empty rows)
    writer = csv.writer(output)
    writer.writerow(['spam', 'eggs', 'bacon', 'ham'])
    output.close()


# parse json
def dl_weather_json(location):
    import json, requests, sys
    url = 'http://api.openweathermap.org/data/2.5/forecast/daily?q={}&cnt=3'.format(location)
    response = requests.get(url)
    response.raise_for_status()
    pval = json.loads(response.text)
    w = pval['list']
    print(w[0]['weather'][0]['main'], '-', w[0]['weather'][0]['description'])
    print(w[1]['weather'][0]['main'], '-', w[1]['weather'][0]['description'])
    print(w[2]['weather'][0]['main'], '-', w[2]['weather'][0]['description'])


# datetime and time examples (pass DT object, ex.: datetime.datetime(2017, 10, 31, 0, 0, 0))
def pause_until(day):
    import datetime, time
    while datetime.datetime.now() < day:
        time.sleep(1)


# multithreading
def take_a_nap(seconds):
    import time
    time.sleep(seconds)
    print('Wake up!')


def two_threads(delay):
    import threading
    print('program started.')
    threado = threading.Thread(target=take_a_nap, args=[delay])
    threado.start()
    print('program stopped.')


# download xkcd archive
def xkcd():
    import requests, os, bs4
    url = 'http://xkcd.com'
    os.makedirs('xkcd', exist_ok=True)
    while not url.endswith('#'):
        res = requests.get(url)
        res.raise_for_status()
        soup = bs4.BeautifulSoup(res.text, "lxml")
        comicel = soup.select('#comic img')
        if comicel == []:
            print('Could not find image')
        else:
            comicurl = 'http:' + comicel[0].get('src')
            print('Downloading image {}'.format(comicurl))
            res = requests.get(comicurl)
            res.raise_for_status()
            imgfile = open(os.path.join('xkcd', os.path.basename(comicurl)), 'wb')
            for chunk in res.iter_content(100000):
                imgfile.write(chunk)
            imgfile.close()
        prevlink = soup.select('a[rel="prev"]')[0]
        url = 'http://xkcd.com' + prevlink.get('href')
    print('Done.')


def download_xkcd(start, end):
    import requests, bs4, os
    url = 'http://xkcd.com/'
    for urlnum in range(start, end):
        res = requests.get('{}{}'.format(url, urlnum))
        res.raise_for_status()
        soup = bs4.BeautifulSoup(res.text, "lxml")
        comicelem = soup.select('#comic img')
        if comicelem == []:
            print('Couldn not find image.')
        else:
            comicurl = comicelem[0].get('src')
            print('Downloading image {}'.format(comicurl))
            res = requests.get(comicurl)
            res.raise_for_status()
            imgfile = open(os.path.join('xkcd', os.path.basename(comicurl)), 'wb')
            for chunk in res.iter_content(100000):
                imgfile.write(chunk)
            imgfile.close()


def xkcd_multi():
    import os, threading
    os.makedirs('xkcd', exist_ok=True)

    downloadthreads = []
    for i in range(0, 1400, 100):
        dlthread = threading.Thread(target=download_xkcd, args=(i, i + 99))
        downloadthreads.append(dlthread)
        dlthread.start()

    for dlthread in downloadthreads:
        dlthread.join()
    print('Done.')


def simple_alarm():
    import time, subprocess
    timeLeft = 60
    while timeLeft > 0:
        print(timeLeft, end='')
        time.sleep(1)
        timeLeft -= 1

    subprocess.Popen(['start', 'alarm.wav'], shell=True)


def send_email(sender, receiver, message):
    import smtplib
    smtpObj = smtplib.SMTP('smtp.gmail.com', 587)  # smtplib.SMTP_SSL('smtp.gmail.com', 465)
    smtpObj.ehlo()
    smtpObj.starttls()
    password = '{}'.format(input('Enter password for {}'.format(sender)))
    smtpObj.login(sender, password)
    status = smtpObj.sendmail(sender, receiver,
                     'Subject:{}\n{}'.format(message[:13], message))
    if status != {}:
        print('There was a problem sending message(s): {}'.format(status))
    smtpObj.quit()


def read_email(receiver):
    import imapclient, imaplib, pprint
    imaplib._MAXLINE = 1000000

    imapObj = imapclient.IMAPClient('imap.gmail.com', ssl=True)
    password = '{}'.format(input('Enter password for {}'.format(receiver)))
    imapObj.login(receiver, password)
    pprint.pprint(imapObj.list_folders())
    imapObj.select_folder('INBOX', readonly=True)
    uids = imapObj.search(['ALL'])  # imapObj.gmail_search('meaning of life')
    rawMessages = imapObj.fetch(uids, ['BODY[]'])
    pprint.pprint(rawMessages)
    # import pyzmail
    # message = pyzmail.PyzMessage.factory(rawMessages[1]['BODY[]'])
    # if (message.text_part != None):
    #     message.text_part.get_payload().decode(message.text_part.charset)
    imapObj.logout()


# Twilio example
def send_text(receiver, message):
    from twilio.rest import TwilioRestClient
    accountSID = ''
    authToken = ''
    client = TwilioRestClient(accountSID, authToken)
    twilioNum = ''
    message = client.messages.create(body=message, from_=twilioNum, to=receiver)


def auto_win_manipulation_test():
    import pyautogui
    pyautogui.PAUSE = 1
    pyautogui.FAILSAFE = True
    for i in range(10):
        pyautogui.moveTo(100, 100, duration=0.25)
        pyautogui.moveTo(200, 100, duration=0.25)
        pyautogui.moveTo(200, 200, duration=0.25)
        pyautogui.moveTo(100, 200, duration=0.25)


def where_is_my_mouse():
    import pyautogui
    print('You can press Ctrl-C at any time to quit.')
    try:
        while True:
            x, y = pyautogui.position()
            positionStr = 'X: {} Y: {}'.format(str(x).rjust(4), str(y).rjust(4))
            pixelColor = pyautogui.screenshot().getpixel((x, y))
            positionStr += ' RGB:({}, {}, {})'.format(str(pixelColor[0]).rjust(3),
                                                      str(pixelColor[1]).rjust(3),
                                                      str(pixelColor[2]).rjust(3))
            print(positionStr, end='')
            print('\b' * len(positionStr), end='', flush=True)
    except KeyboardInterrupt:
        print('\nQuit.')


def draw_spiral_mouse():
    import pyautogui, time
    time.sleep(5)
    pyautogui.click()
    distance = 200
    while distance > 0:
        pyautogui.dragRel(distance, 0, duration=0.2)
        distance -= 5
        pyautogui.dragRel(0, distance, duration=0.2)
        pyautogui.dragRel(-distance, 0, duration=0.2)
        distance -= 5
        pyautogui.dragRel(0, -distance, duration=0.2)


def recognize_image(file):
    import pyautogui
    pos = pyautogui.locateOnScreen(file)
    print(str(pos))
